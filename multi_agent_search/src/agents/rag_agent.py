"""
RAG Agent with LangChain - Document Processing & Semantic Search
Based on: https://docs.langchain.com/oss/python/langchain/rag

Features:
- PDF, DOCX, Audio, Image processing
- Vector store (Supabase pgvector)
- Semantic search with embeddings
- LangGraph agent with retrieval tool
"""

import os
import logging
from typing import List, Optional, Literal, Union
from pathlib import Path

from langchain.chat_models import init_chat_model
from langchain.tools import tool
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document
from langgraph.prebuilt import create_react_agent
from langchain_ollama import OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
try:
    from langchain_community.document_loaders import (
        PyPDFLoader,
        Docx2txtLoader,
        CSVLoader,
    )
except ImportError:
    # Fallback: manual loaders
    PyPDFLoader = None
    Docx2txtLoader = None
    CSVLoader = None

import pymupdf4llm  # High-quality PDF to Markdown

from ..config import settings

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ============ LANGSMITH ============
def setup_langsmith():
    """LangSmith tracing for RAG agent"""
    tracing_enabled = os.getenv("LANGSMITH_TRACING", "").strip().lower() in {"1", "true", "yes", "on"}
    if not tracing_enabled:
        return False
    if not settings.langsmith_api_key:
        return False

    os.environ["LANGSMITH_TRACING"] = "true"
    os.environ["LANGSMITH_API_KEY"] = settings.langsmith_api_key

    # Standard LangChain vars
    os.environ["LANGCHAIN_TRACING_V2"] = "true"
    os.environ["LANGCHAIN_API_KEY"] = settings.langsmith_api_key
    os.environ["LANGCHAIN_PROJECT"] = os.getenv("LANGCHAIN_PROJECT", "ai-research-rag-agent")

    logger.info("[LANGSMITH] Tracing aktif")
    return True


# ============ VECTOR STORE SETUP ============

# Initialize embeddings (Ollama - Local & Free)
# Override via env: OLLAMA_EMBED_MODEL
_embeddings = OllamaEmbeddings(
    model=os.getenv("OLLAMA_EMBED_MODEL", "nomic-embed-text:latest"),
    base_url=settings.ollama_base_url,
    num_gpu=-1,
)

# ChromaDB persistent directory
_chroma_dir = Path(settings.project_root) / "chroma_db"
_chroma_dir.mkdir(parents=True, exist_ok=True)

# Initialize ChromaDB vector store (persistent)
_vector_store = Chroma(
    collection_name="rag_documents",
    embedding_function=_embeddings,
    persist_directory=str(_chroma_dir),
)

logger.info(f"[RAG] ChromaDB initialized at: {_chroma_dir}")

# Track ingested sources to prevent duplicates
_ingested_sources = set()


def _uploads_dir() -> Path:
    # Normalize to project uploads directory
    return Path(settings.project_root) / "uploads"


_INGESTED_SOURCES_FILE = _uploads_dir() / ".rag_ingested_sources.json"


def _load_ingested_sources() -> None:
    """Load ingested sources from disk (and optionally Supabase if configured).

    This is best-effort; it should never crash import.
    """
    global _ingested_sources

    # 1) Disk persistence (always available)
    try:
        if _INGESTED_SOURCES_FILE.exists():
            import json
            data = json.loads(_INGESTED_SOURCES_FILE.read_text(encoding="utf-8"))
            if isinstance(data, list):
                _ingested_sources.update(str(x) for x in data if str(x).strip())
    except Exception as e:
        logger.warning(f"[RAG] Failed to load ingested sources from disk: {e}")

    # 2) Supabase persistence (optional)
    # Requires a table named `rag_sources` with at least:
    # - source_name text primary key
    # - created_at timestamptz default now()
    try:
        if settings.supabase_url and settings.supabase_key:
            try:
                from supabase import create_client  # type: ignore
            except Exception:
                create_client = None
            if create_client:
                client = create_client(settings.supabase_url, settings.supabase_key)
                resp = client.table("rag_sources").select("source_name").limit(5000).execute()
                for row in resp.data or []:
                    name = row.get("source_name")
                    if name:
                        _ingested_sources.add(str(name))
    except Exception as e:
        # Don't fail if table doesn't exist or credentials are limited.
        logger.info(f"[RAG] Supabase rag_sources not loaded (optional): {e}")


def _persist_ingested_sources() -> None:
    """Persist ingested sources to disk and (optionally) Supabase."""
    # 1) Disk
    try:
        import json
        _uploads_dir().mkdir(parents=True, exist_ok=True)
        _INGESTED_SOURCES_FILE.write_text(
            json.dumps(sorted(_ingested_sources), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception as e:
        logger.warning(f"[RAG] Failed to persist ingested sources to disk: {e}")

    # 2) Supabase (best-effort)
    try:
        if settings.supabase_url and settings.supabase_key:
            try:
                from supabase import create_client  # type: ignore
            except Exception:
                create_client = None
            if create_client:
                client = create_client(settings.supabase_url, settings.supabase_key)
                # Upsert all sources (requires primary key on source_name)
                rows = [{"source_name": s} for s in _ingested_sources]
                if rows:
                    client.table("rag_sources").upsert(rows).execute()
    except Exception as e:
        logger.info(f"[RAG] Supabase rag_sources not persisted (optional): {e}")

# ============ GLOBAL CHUNKING STRATEGY (ROBUST) ============
# Goal (global): keep "image + its local explanation" together, and avoid micro-chunks.
# Why not SemanticChunker here:
# - For PDF→Markdown, it often splits into single-sentence chunks (bad retrieval and bad answers)
# - It can cut across headings/images, mixing sections
#
# Strategy:
# 1) If content looks like Markdown (headings/images), split by headings up to H4.
#    (We intentionally DO NOT hard-split on H5/H6 to avoid separating algorithm heading from its description.)
# 2) Inside each section, split by images so each image gets its own chunk when multiple exist.
# 3) If a chunk is still too large, fall back to a gentle recursive splitter.

import re


def _looks_like_markdown(text: str) -> bool:
    if not text:
        return False
    # headings or markdown images are strong signals
    return ("\n#" in text) or ("![](" in text) or bool(re.search(r"(?m)^#{1,6}\s+", text))


_fallback_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1200,
    chunk_overlap=250,  # ✅ GELİŞME 4: Arttırıldı (120→250) - Görseller context'ten kopmasın
    add_start_index=True,
    separators=[
        "\n\n",  # paragraphs
        "\n",    # lines
        " ",
        "",
    ],
    is_separator_regex=False,
)


def _split_markdown_by_headings(md: str, max_heading_level: int = 4) -> List[str]:
    """Split markdown into sections using headings up to max_heading_level.
    
    ✅ FIX: Split BEFORE H4 (level < 4), not at H4
    This ensures H4 headers stay with their content
    """
    md = (md or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = md.split("\n")
    sections: List[List[str]] = []
    current: List[str] = []

    heading_re = re.compile(r"^(#{1,6})\s+.*")
    for line in lines:
        m = heading_re.match(line.strip())
        if m:
            level = len(m.group(1))
            # ✅ Split BEFORE H4 (consistent with agentic_chunker.py)
            if level < max_heading_level and current:
                sections.append(current)
                current = [line]
                continue
        current.append(line)

    if current:
        sections.append(current)

    return ["\n".join(s).strip() for s in sections if "\n".join(s).strip()]


def _split_section_by_images(section_text: str) -> List[str]:
    """Within a section, split into smaller chunks by image boundaries.

    If there are multiple images in the same section (e.g., conclusion page),
    each image becomes its own chunk along with nearby text.
    """
    section_text = (section_text or "").strip()
    if not section_text:
        return []

    lines = section_text.split("\n")
    image_line_re = re.compile(r"^!\[[^\]]*\]\([^)]+\)\s*$")

    chunks: List[str] = []
    current: List[str] = []
    current_has_image = False
    recent_nonempty: List[str] = []

    for line in lines:
        is_image = bool(image_line_re.match(line.strip()))

        if is_image and current_has_image:
            # finalize previous image-chunk
            chunk = "\n".join(current).strip()
            if chunk:
                chunks.append(chunk)
            current = []
            current_has_image = False

        if is_image and not current:
            # pull a little context before the image, if available
            for ctx in recent_nonempty[-6:]:
                current.append(ctx)

        current.append(line)
        if is_image:
            current_has_image = True

        if line.strip():
            recent_nonempty.append(line)

    tail = "\n".join(current).strip()
    if tail:
        chunks.append(tail)

    return chunks


def _make_rag_splits(text: str, source_name: str) -> List[Document]:
    """Create RAG-ready Document splits with stable semantics."""
    base_doc = Document(page_content=text or "", metadata={"source": source_name})

    if not _looks_like_markdown(base_doc.page_content):
        splits = _fallback_splitter.split_documents([base_doc])
        for s in splits:
            s.metadata["source"] = source_name
            s.metadata["has_images"] = False
        return splits

    # Markdown path
    heading_sections = _split_markdown_by_headings(base_doc.page_content, max_heading_level=4)

    chunks: List[Document] = []
    for sec in heading_sections:
        # If section has multiple images, split them into separate chunks.
        subchunks = _split_section_by_images(sec) if "![](" in sec or "![" in sec else [sec]

        for sub in subchunks:
            if not sub.strip():
                continue

            doc = Document(page_content=sub.strip(), metadata={"source": source_name})

            # If a chunk is still too big (rare), gently split by paragraphs.
            if len(doc.page_content) > 2000:
                for s in _fallback_splitter.split_documents([doc]):
                    s.metadata["source"] = source_name
                    s.metadata["has_images"] = bool(re.search(r"!\[[^\]]*\]\([^)]+\)", s.page_content))
                    chunks.append(s)
            else:
                doc.metadata["has_images"] = bool(re.search(r"!\[[^\]]*\]\([^)]+\)", doc.page_content))
                chunks.append(doc)

    return chunks


def _safe_resolve_upload_path(img_path: str) -> Optional[Path]:
    """Resolve a markdown image path safely under uploads/.

    Prevents path traversal (e.g. ../../secret) and disallows absolute paths.
    Returns a resolved absolute Path if valid and exists, else None.
    """
    if not img_path:
        return None

    # Disallow URLs and absolute paths
    lowered = img_path.strip().lower()
    if lowered.startswith("http://") or lowered.startswith("https://"):
        return None

    # 🔒 SECURITY: Block path traversal attempts
    if ".." in img_path or "~" in img_path:
        logger.warning(f"[SECURITY] Path traversal blocked: {img_path}")
        return None

    p = Path(img_path)
    if p.is_absolute():
        return None

    # Fix: If path already starts with 'uploads/', strip it to avoid duplication
    # e.g. "uploads/image.png" -> "image.png" because we join with uploads_dir
    parts = list(p.parts)
    if parts and parts[0] == "uploads":
        p = Path(*parts[1:])

    uploads = _uploads_dir().resolve()
    resolved = (uploads / p).resolve()
    
    # 🔒 SECURITY: Verify resolved path is under uploads/
    try:
        resolved.relative_to(uploads)
    except ValueError:
        logger.warning(f"[SECURITY] Path outside uploads blocked: {img_path}")
        return None

    if not resolved.exists():
        return None

    return resolved


# ============ DOCUMENT PROCESSING TOOLS ============

@tool
def load_pdf(file_path: str) -> str:
    """
    Load and extract text from a PDF file with images using PyMuPDF4LLM.
    Extracts text, tables, and saves page images.
    
    Args:
        file_path: Absolute path to PDF file
        
    Returns:
        Extracted text content (Markdown)
    """
    # 1. Try PyMuPDF4LLM (Fast, extracts images, good Markdown)
    try:
        import pymupdf4llm
        
        logger.info(f"[PDF] Processing {Path(file_path).name} with PyMuPDF4LLM...")
        
        # Create images folder
        images_folder = Path(file_path).parent / f"{Path(file_path).stem}_images"
        images_folder.mkdir(exist_ok=True)
        
        # Convert PDF to Markdown with images
        md_text = pymupdf4llm.to_markdown(
            file_path,
            write_images=True,
            image_path=str(images_folder),
            image_format="png"
        )
        
        logger.info(f"[PDF] ✅ Converted {len(md_text)} chars with images (PyMuPDF4LLM)")
        
        # Save debug file
        debug_path = Path(file_path).parent / f"{Path(file_path).stem}_pymupdf_debug.md"
        debug_path.write_text(md_text, encoding="utf-8")
        logger.info(f"[PDF] 📝 Markdown saved to: {debug_path}")
        logger.info(f"[PDF] 🖼️ Images saved to: {images_folder}")
        
        if len(md_text) >= 100:
            return md_text
            
    except Exception as e:
        logger.warning(f"[PDF] PyMuPDF4LLM failed ({e}), falling back to Unstructured...")
    
    # 2. Fallback: Unstructured
    try:
        from unstructured.partition.pdf import partition_pdf
        
        logger.info(f"[PDF] Processing {Path(file_path).name} with Unstructured (fallback)...")
        
        elements = partition_pdf(filename=file_path, strategy="fast")
        
        text_parts = []
        for el in elements:
            element_text = str(el)
            element_type = el.category if hasattr(el, 'category') else 'Unknown'
            
            if element_type == 'Title':
                text_parts.append(f"# {element_text}")
            elif element_type == 'Header':
                text_parts.append(f"## {element_text}")
            elif element_type == 'Table':
                text_parts.append(f"\n{element_text}\n")
            else:
                text_parts.append(element_text)
        
        text = "\n\n".join(text_parts)
        logger.info(f"[PDF] ✅ Converted {len(text)} chars from {len(elements)} elements (Unstructured OCR)")
        
        debug_path = Path(file_path).parent / f"{Path(file_path).stem}_unstructured_debug.md"
        debug_path.write_text(text, encoding="utf-8")
        logger.info(f"[PDF] 📝 Unstructured output saved to: {debug_path}")
        
        if len(text) >= 50:
            return text
            
    except Exception as e2:
        logger.warning(f"[PDF] Unstructured failed ({e2}), falling back to PyPDF...")

    # 3. Last Fallback: PyPDF (Simple text extraction)
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = "\n\n".join([page.extract_text() for page in reader.pages])
        logger.info(f"[PDF] Loaded {len(reader.pages)} pages, {len(text)} chars (PyPDF - last resort)")
        
        debug_path = Path(file_path).parent / f"{Path(file_path).stem}_pypdf_debug.txt"
        debug_path.write_text(text, encoding="utf-8")
        logger.info(f"[PDF] 📝 PyPDF output saved to: {debug_path}")
        
        return text
    except Exception as e3:
        logger.error(f"[PDF] All methods failed: {e3}")
        return f"Error loading PDF: {str(e3)}"


@tool
def load_docx(file_path: str) -> str:
    """
    Load and extract text + images from a DOCX file using python-docx.
    Extracts text, tables, and saves embedded images similar to PDF processing.
    
    Args:
        file_path: Absolute path to DOCX file
        
    Returns:
        Extracted text content (Markdown with image references)
    """
    # 1. Try python-docx with image extraction (Primary method)
    try:
        from docx import Document as DocxDocument
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        import io
        
        logger.info(f"[DOCX] Processing {Path(file_path).name} with python-docx...")
        
        # Create images folder
        images_folder = Path(file_path).parent / f"{Path(file_path).stem}_images"
        images_folder.mkdir(exist_ok=True)
        
        doc = DocxDocument(file_path)
        markdown_parts = []
        image_counter = 0
        
        # Extract images from document
        image_map = {}  # Map relationship ID to saved filename
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    # Determine extension from content type
                    ext_map = {
                        'image/png': 'png',
                        'image/jpeg': 'jpg',
                        'image/jpg': 'jpg',
                        'image/gif': 'gif',
                        'image/bmp': 'bmp'
                    }
                    ext = ext_map.get(rel.target_part.content_type, 'png')
                    
                    # Save image
                    image_filename = f"{Path(file_path).stem}-{image_counter}.{ext}"
                    image_path = images_folder / image_filename
                    image_path.write_bytes(image_data)
                    
                    # Store relative path for markdown
                    relative_path = f"{images_folder.name}/{image_filename}"
                    image_map[rel_id] = relative_path
                    image_counter += 1
                    
                except Exception as e:
                    logger.warning(f"[DOCX] Failed to extract image {rel_id}: {e}")
        
        # Process paragraphs and tables in order
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                text = para.text.strip()
                
                # Check paragraph style for headers
                style_name = para.style.name if para.style else ""
                
                if style_name.startswith('Heading 1'):
                    markdown_parts.append(f"# {text}")
                elif style_name.startswith('Heading 2'):
                    markdown_parts.append(f"## {text}")
                elif style_name.startswith('Heading 3'):
                    markdown_parts.append(f"### {text}")
                elif style_name.startswith('Heading 4'):
                    markdown_parts.append(f"#### {text}")
                elif text:
                    markdown_parts.append(text)
                
                # Check for inline images
                for run in para.runs:
                    for rel_id in run.element.xpath('.//a:blip/@r:embed', namespaces={
                        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                    }):
                        if rel_id in image_map:
                            markdown_parts.append(f"![Image]({image_map[rel_id]})")
                
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                # Convert table to markdown
                markdown_parts.append("\n")
                for i, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    markdown_parts.append("| " + " | ".join(cells) + " |")
                    if i == 0:  # Add separator after header
                        markdown_parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
                markdown_parts.append("\n")
        
        md_text = "\n\n".join(markdown_parts)
        logger.info(f"[DOCX] ✅ Converted {len(md_text)} chars with {image_counter} images (python-docx)")
        
        # Save debug file
        debug_path = Path(file_path).parent / f"{Path(file_path).stem}_docx_debug.md"
        debug_path.write_text(md_text, encoding="utf-8")
        logger.info(f"[DOCX] 📝 Markdown saved to: {debug_path}")
        if image_counter > 0:
            logger.info(f"[DOCX] 🖼️ {image_counter} images saved to: {images_folder}")
        
        if len(md_text) >= 50:
            return md_text
            
    except Exception as e:
        logger.warning(f"[DOCX] python-docx failed ({e}), falling back to Unstructured...")
    
    # 2. Fallback: Unstructured (no image extraction)
    try:
        from unstructured.partition.docx import partition_docx
        
        logger.info(f"[DOCX] Processing {Path(file_path).name} with Unstructured (fallback)...")
        
        elements = partition_docx(filename=file_path)
        text = "\n\n".join([str(el) for el in elements])
        
        logger.info(f"[DOCX] ✅ Converted {len(text)} chars (Unstructured)")
        
        debug_path = Path(file_path).parent / f"{Path(file_path).stem}_unstructured_debug.md"
        debug_path.write_text(text, encoding="utf-8")
        logger.info(f"[DOCX] 📝 Markdown saved to: {debug_path}")
        
        return text
        
    except Exception as e2:
        logger.warning(f"[DOCX] Unstructured failed ({e2}), falling back to Docx2txt...")
        
        # 3. Last Fallback: Docx2txt (text only)
        try:
            if Docx2txtLoader:
                loader = Docx2txtLoader(file_path)
                docs = loader.load()
                text = "\n\n".join([doc.page_content for doc in docs])
                return text
            else:
                import docx2txt
                return docx2txt.process(file_path)
        except Exception as e3:
            logger.error(f"[DOCX] All methods failed: {e3}")
            return f"Error loading DOCX: {str(e)}"


@tool
def load_csv(file_path: str) -> str:
    """
    Load and extract data from a CSV file.
    
    Args:
        file_path: Absolute path to CSV file
        
    Returns:
        CSV data as formatted text
    """
    try:
        if CSVLoader:
            loader = CSVLoader(file_path)
            docs = loader.load()
            text = "\n\n".join([doc.page_content for doc in docs])
            logger.info(f"[CSV] Loaded {len(docs)} rows from {Path(file_path).name}")
            return text
        else:
            # Fallback: csv module
            import csv
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                rows = [str(dict(row)) for row in reader]
                text = "\n\n".join(rows)
                logger.info(f"[CSV] Loaded {len(rows)} rows from {Path(file_path).name}")
                return text
    except Exception as e:
        logger.error(f"[CSV] Error loading {file_path}: {e}")
        return f"Error loading CSV: {str(e)}"


@tool
async def transcribe_audio(file_path: str, language: str = "tr") -> str:
    """
    Transcribe audio file locally.
    
    Args:
        file_path: Absolute path to audio file (mp3, wav, m4a)
        language: Language code (default: tr for Turkish)
        
    Returns:
        Transcribed text
    """
    try:
        # Local STT recommendation: faster-whisper
        # Install (optional): pip install faster-whisper ffmpeg-python
        try:
            from faster_whisper import WhisperModel  # type: ignore
        except Exception:
            return (
                "Local STT için `faster-whisper` önerilir.\n"
                "Kurulum: `pip install faster-whisper ffmpeg-python` ve sistemde ffmpeg olmalı.\n"
                "Alternatif: whisper.cpp (CLI)."
            )

        model_name = os.getenv("WHISPER_MODEL", "small")
        device = os.getenv("WHISPER_DEVICE", "cpu")
        compute_type = os.getenv("WHISPER_COMPUTE_TYPE", "int8")

        whisper_model = WhisperModel(model_name, device=device, compute_type=compute_type)
        segments, info = whisper_model.transcribe(file_path, language=language)
        text = "".join(seg.text for seg in segments).strip()
        logger.info(f"[AUDIO] Transcribed {Path(file_path).name} ({info.language}) - {len(text)} chars")
        return text or "(Boş transkripsiyon)"
    except Exception as e:
        logger.error(f"[AUDIO] Error transcribing {file_path}: {e}")
        return f"Error transcribing audio: {str(e)}"


@tool
async def analyze_image(file_path: str, prompt: str = "Bu görseli detaylı Türkçe açıkla") -> str:
    """
    Analyze image locally via Ollama vision model.
    
    Args:
        file_path: Absolute path to image file
        prompt: Analysis prompt (default: Turkish description)
        
    Returns:
        Image analysis text
    """
    try:
        import base64
        import httpx

        with open(file_path, "rb") as image_file:
            image_b64 = base64.b64encode(image_file.read()).decode("utf-8")

        vision_model = os.getenv("OLLAMA_VISION_MODEL", "llava:latest")

        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                f"{settings.ollama_base_url}/api/chat",
                json={
                    "model": vision_model,
                    "messages": [
                        {"role": "user", "content": prompt, "images": [image_b64]}
                    ],
                    "stream": False,
                    "options": {"num_gpu": -1},
                },
            )

        if response.status_code != 200:
            return f"Ollama Vision error: HTTP {response.status_code} - {response.text}"

        data = response.json()
        text = (data.get("message", {}) or {}).get("content", "")
        logger.info(f"[IMAGE] Analyzed {Path(file_path).name} with {vision_model}")
        return text or "(Boş analiz)"
    except Exception as e:
        logger.error(f"[IMAGE] Error analyzing {file_path}: {e}")
        return f"Error analyzing image: {str(e)}"


def validate_chunk_quality(chunks: List[Document]) -> List[Document]:
    """
    Filter out low-quality chunks that hurt retrieval.
    
    Bad chunks:
    - Too short (< 50 chars)
    - Only whitespace/formatting
    - Orphaned image references (image link but no text explanation nearby)
    """
    validated = []
    import re
    
    for chunk in chunks:
        content = chunk.page_content.strip()
        
        # Filter 1: Too short (likely formatting noise)
        if len(content) < 50 and not chunk.metadata.get("has_images", False):
            continue
        
        # Filter 2: Only whitespace/formatting
        text_only = re.sub(r'[#*\-_\[\]()!\n\s]+', '', content)
        if len(text_only) < 15 and not chunk.metadata.get("has_images", False):
            continue
        
        # Filter 3: Orphaned images (image with < 20 chars of explanatory text)
        if "![" in content:
            text_without_images = re.sub(r'!\[[^\]]*\]\([^)]+\)', '', content)
            # If there's almost no text, but we HAVE an image, we keep it 
            # only if it's potentially important (Always-On vision will handle it)
            if len(text_without_images.strip()) < 10 and len(content) > 100:
                # Likely just a large image link, keep it but log it
                pass
        
        validated.append(chunk)
    
    if len(validated) < len(chunks):
        logger.info(f"[VALIDATE] Kept {len(validated)}/{len(chunks)} quality chunks")
    
    return validated


# ============ RAG CORE FUNCTIONS ============

def ingest_documents(documents: List[Document], source_name: str = "upload") -> int:
    """
    Ingest documents into vector store.
    """
    # Split documents into chunks
    all_splits: List[Document] = []
    for doc in documents:
        all_splits.extend(_make_rag_splits(doc.page_content, source_name))
    
    # ✅ VALIDATE QUALITY
    all_splits = validate_chunk_quality(all_splits)
    
    # Add source metadata
    for split in all_splits:
        split.metadata["source"] = source_name
    
    # Add to vector store
    if all_splits:
        _vector_store.add_documents(documents=all_splits)
        _ingested_sources.add(source_name)
        _persist_ingested_sources()
    
    logger.info(f"[RAG] Ingested {len(all_splits)} chunks from {source_name}")
    return len(all_splits)


def ingest_text(text: str, source_name: str = "upload") -> int:
    """
    Ingest raw text into vector store with intelligent chunking.
    """
    # Check if already ingested
    if source_name in _ingested_sources:
        logger.warning(f"[RAG] ⚠️ Duplicate: {source_name} already ingested, skipping...")
        return 0
    
    # Choose chunking strategy
    use_agentic = os.getenv("USE_AGENTIC_CHUNKER", "true").lower() in ("1", "true", "yes", "on")
    
    if use_agentic:
        try:
            from .agentic_chunker import agentic_chunk_text
            logger.info(f"[RAG] 🤖 Using Agentic Chunker (qwen2.5:7b) for: {source_name}")
            final_splits = agentic_chunk_text(text, source_name)
        except Exception as e:
            logger.warning(f"[RAG] Agentic chunking failed ({e}), falling back to markdown splitter")
            final_splits = _make_rag_splits(text, source_name)
    else:
        final_splits = _make_rag_splits(text, source_name)

    # ✅ VALIDATE QUALITY
    final_splits = validate_chunk_quality(final_splits)

    # Add source metadata
    for split in final_splits:
        split.metadata["source"] = source_name

    # Add to vector store
    if final_splits:
        _vector_store.add_documents(documents=final_splits)
        _ingested_sources.add(source_name)
        _persist_ingested_sources()
    
    logger.info(f"[RAG] ✅ Ingested {len(final_splits)} chunks from {source_name}")
    return len(final_splits)



# ============ RETRIEVAL TOOL ============

from pydantic import BaseModel, Field

class RetrieveContextInput(BaseModel):
    """Input schema for retrieve_context tool"""
    query: str = Field(description="Search query to find relevant document chunks")
    # Groq API strict validation fix: Accept string, convert to int internally
    top_k: str = Field(default="3", description="Number of top results to return (e.g. '3')")

@tool(args_schema=RetrieveContextInput, response_format="content_and_artifact")
def retrieve_context(query: str, top_k: str = "3"):
    """
    Retrieve relevant context from ingested documents using semantic search.
    AUTOMATICALLY analyzes images found in the context using vision models.
    
    Args:
        query: Search query
        top_k: Number of top results to return (default: "3" - balanced for speed)
        
    Returns:
        Retrieved context as formatted text with source metadata AND image analysis
    """
    # Ensure top_k is an integer
    try:
        k_val = int(top_k)
    except (ValueError, TypeError):
        k_val = 3  # Balanced for speed and coverage

    # ENSEMBLE RETRIEVAL: BM25 (keyword) + Semantic (context)
    # This works for ANY PDF, not just one specific document
    
    # 1. BM25 Retrieval (keyword-based, excellent for Turkish)
    from rank_bm25 import BM25Okapi
    import re
    
    # Get all documents for BM25
    all_docs_data = _vector_store.get()
    if not all_docs_data or 'documents' not in all_docs_data:
        return "No relevant documents found.", []
    
    all_docs_content = all_docs_data['documents']
    all_docs_metadata = all_docs_data['metadatas']
    
    # ✅ Turkish-aware tokenization for BM25
    def tokenize(text):
        # Turkish lowercase mapping (i→i, I→ı, İ→i, Ş→ş, Ğ→ğ, etc.)
        text = text.replace('I', 'ı').replace('İ', 'i').lower()
        
        # Extract tokens
        tokens = re.findall(r'\w+', text)
        
        # Remove Turkish stopwords for better relevance
        turkish_stopwords = {
            "bir", "bu", "da", "de", "ve", "ile", "için", "mi", "mı", 
            "mu", "mü", "gibi", "daha", "çok", "en", "var", "yok",
            "olan", "olarak", "ama", "fakat", "veya", "hem", "ya", "ne"
        }
        tokens = [t for t in tokens if t not in turkish_stopwords and len(t) > 1]
        
        return tokens
    
    tokenized_corpus = [tokenize(doc) for doc in all_docs_content]
    bm25 = BM25Okapi(tokenized_corpus)
    
    # BM25 search
    tokenized_query = tokenize(query)
    bm25_scores = bm25.get_scores(tokenized_query)
    
    # 2. Semantic search
    semantic_docs = _vector_store.similarity_search_with_score(query, k=len(all_docs_content))
    semantic_scores = {i: 1 / (1 + score) for i, (_, score) in enumerate(semantic_docs)}  # Lower distance = higher score
    
    # 3. Ensemble: BM25 (0.6) + Semantic (0.4)
    ensemble_scores = []
    for i in range(len(all_docs_content)):
        # Normalize BM25 scores to 0-1
        bm25_normalized = bm25_scores[i] / (max(bm25_scores) + 1e-6)
        semantic_normalized = semantic_scores.get(i, 0)
        
        # Weighted ensemble
        final_score = 0.6 * bm25_normalized + 0.4 * semantic_normalized
        
        # ✅ GELİŞME 1: DİNAMİK GÖRSEL NİYETİ BOOST (Generic RAG için kritik)
        query_keywords = set(tokenize(query))
        visual_keywords = {
            # Türkçe görsel kelimeleri
            "görsel", "göster", "resim", "foto", "fotoğraf", "şekil", "figür",
            "grafik", "chart", "çizim", "diyagram", "diagram",
            "tablo", "table", "matris", "matrix",
            # Soru kelimeleri
            "var", "varmı", "varmi", "nerede", "göster", "bul", "getir",
            # İngilizce
            "image", "picture", "photo", "figure", "graph", "plot", "visualization"
        }
        has_visual_query = bool(query_keywords & visual_keywords)
        has_images = all_docs_metadata[i].get("has_images", False)
        
        # ✅ SMART BOOST: Kullanıcı niyetine göre dinamik ağırlıklandırma
        if has_visual_query and has_images:
            # Kullanıcı görsel arıyor VE chunk'ta görsel var → %50 boost
            final_score *= 1.5
        elif not has_visual_query and has_images:
            # Kullanıcı görsel aramIyor AMA chunk'ta görsel var → hafif ceza (metne odaklan)
            final_score *= 0.9
        # Else: Görsel yok veya neutral durum → skor değişmez
        
        ensemble_scores.append((final_score, i))
    
    # Sort by ensemble score and take top k
    ensemble_scores.sort(key=lambda x: x[0], reverse=True)
    
    # ✅ SMART FILTER: İlgili chunk'ları bul
    if ensemble_scores:
        top_score = ensemble_scores[0][0]
        score_threshold = max(top_score * 0.7, 0.3)
        
        # ✅ YENİ: Chunk başlığına/içeriğine göre filtrele
        # "logistic regression ile alakalı görsel" → Logistic Regression CHUNK'ı, karşılaştırma değil
        content_keywords = query_keywords - visual_keywords
        
        filtered_scores = []
        for score, idx in ensemble_scores[:k_val]:
            if score < score_threshold:
                continue
            
            # Chunk metadata ve content'i kontrol et
            chunk_meta = all_docs_metadata[idx]
            chunk_content = all_docs_content[idx]
            
            # Başlık varsa kontrol et
            chunk_title = chunk_meta.get("title", "").lower()
            
            # ✅ KRİTİK CHECK: Query'deki ana keyword chunk başlığında var mı?
            # Örnek: "logistic regression" query'de → "Logistic Regression" başlıklı chunk ✅
            #        "logistic regression" query'de → "Karmaşıklık Matrisi" başlıklı chunk ❌
            if content_keywords:
                title_tokens = set(tokenize(chunk_title))
                content_tokens = set(tokenize(chunk_content[:500]))  # İlk 500 char
                
                # Başlık match'i yüksek öncelikli
                title_match_count = len(content_keywords & title_tokens)
                content_match_count = len(content_keywords & content_tokens)
                
                # Eğer başlık match yoksa ve content match düşükse, skip
                if title_match_count == 0 and content_match_count < len(content_keywords) * 0.5:
                    logger.info(f"[RAG FILTER] Skipping chunk #{idx} - low relevance (title: {chunk_title[:50]})")
                    continue
            
            filtered_scores.append((score, idx))
        
        if not filtered_scores:
            # If nothing passes threshold, take only the top 1
            filtered_scores = [ensemble_scores[0]]
        
        top_indices = [idx for _, idx in filtered_scores]
    else:
        top_indices = []
    
    # Reconstruct documents
    retrieved_docs = []
    for idx in top_indices:
        doc = Document(
            page_content=all_docs_content[idx],
            metadata=all_docs_metadata[idx]
        )
        retrieved_docs.append(doc)
    
    logger.info(f"[RAG ENSEMBLE] Top chunk scores (BM25+Semantic): {[(round(s, 3), all_docs_metadata[i].get('title', 'N/A')[:50]) for s, i in ensemble_scores[:k_val]]}")
    logger.info(f"[RAG QUALITY] Filtered to {len(retrieved_docs)} chunks (threshold: {score_threshold:.3f})")
    
    if not retrieved_docs:
        return "No relevant documents found.", []
    
    # Extract image paths from markdown content
    import re
    image_blocks = []
    
    base_url = settings.backend_public_url.rstrip("/")
    
    # ============ VISUAL QUERY DETECTION (MUST BE BEFORE VISION ANALYSIS) ============
    # Define visual keywords and tokenizer for filtering
    def tokenize_for_visual(text):
        """Simple tokenizer for visual detection - NO stopwords removal"""
        text = text.replace('I', 'ı').replace('İ', 'i').lower()
        return set(re.findall(r'\w+', text))
    
    query_keywords = tokenize_for_visual(query)
    
    visual_keywords = {
        "görsel", "göster", "resim", "foto", "fotoğraf", "şekil", "figür",
        "grafik", "chart", "çizim", "diyagram", "diagram",
        "tablo", "table", "matris", "matrix", "çizimi", "görüntüsü",
        "var", "varmı", "varmi", "nerede", "hangi", "bul", "getir", "paylaş",
        "image", "picture", "photo", "figure", "graph", "plot", "visualization", "show"
    }
    
    has_visual_query = bool(query_keywords & visual_keywords)
    
    if has_visual_query:
        logger.info(f"[RAG] 🖼️ Visual query detected: {query_keywords & visual_keywords}")
    else:
        logger.info(f"[RAG] 📕 Text-only query detected")
    
    # ✅ GELİŞME 2: AKILLI GÖRSEL PUANLAMA (TÜM CHUNK'LARDA)
    # Generic RAG için kritik: Sadece ilk chunk'a değil, TÜM chunk'lardaki TÜM görselleri puanlayıp en iyiyi seç
    
    def score_image_relevance(query_keywords, visual_keywords, alt_text, img_path, surrounding_text):
        """Generic image scoring algorithm for any document type"""
        score = 0
        content_keywords = query_keywords - visual_keywords  # "ridge", "classifier"
        
        img_metadata = (alt_text + " " + img_path).lower()
        
        for kw in content_keywords:
            # 1. Alt text veya dosya isminde geçiyorsa (En yüksek sinyal)
            if kw in img_metadata:
                score += 10
            # 2. Resmin hemen yakınındaki metinde geçiyorsa (Güçlü sinyal)
            if kw in surrounding_text.lower():
                score += 3
        
        return score
    
    # TÜM chunk'lardaki TÜM görselleri topla ve puanla
    all_potential_images = []
    
    logger.info(f"[RAG] 🔍 Scanning ALL chunks for images (smart multi-chunk analysis)")
    
    for idx, doc in enumerate(retrieved_docs[:5]):  # Top 5 chunk'a bak (balance: quality vs speed)
        img_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        img_matches = re.findall(img_pattern, doc.page_content)
        
        for alt_text, img_path in img_matches:
            safe_path = _safe_resolve_upload_path(img_path)
            if not safe_path:
                continue
            
            # Görselin etrafındaki 300 karakteri al (context window)
            img_position = doc.page_content.find(img_path)
            if img_position == -1:
                img_position = doc.page_content.find(f"![{alt_text}]({img_path})")
            
            if img_position != -1:
                context_start = max(0, img_position - 300)
                context_end = min(len(doc.page_content), img_position + 300)
                surrounding_text = doc.page_content[context_start:context_end]
            else:
                surrounding_text = doc.page_content[:600]  # Fallback: ilk 600 char
            
            # ✅ PUANLAMA: Generic algorithm
            relevance_score = score_image_relevance(
                query_keywords, visual_keywords, alt_text, img_path, surrounding_text
            )
            
            all_potential_images.append({
                "alt": alt_text,
                "path": img_path,
                "safe_path": safe_path,
                "score": relevance_score,
                "chunk_idx": idx,
                "source": doc.metadata.get("source", "unknown")
            })
            
            logger.info(f"[RAG SCORER] Image: {safe_path.name[:50]} | Score: {relevance_score} | Chunk: #{idx}")
    
    # En yüksek skorlu 1-2 görseli seç (Generic: her dökümanda en iyiyi bul)
    all_potential_images.sort(key=lambda x: x['score'], reverse=True)
    best_images = [img for img in all_potential_images if img['score'] > 0][:2]  # Top 2 (or top 1 if score=0)
    
    if not best_images and all_potential_images:
        # Eğer hiçbir keyword match yoksa bile, ilk görseli al (Always-On strategy)
        best_images = [all_potential_images[0]]
    
    logger.info(f"[RAG] ✅ Selected {len(best_images)} best images from {len(all_potential_images)} candidates")
    
    # ✅ GELİŞME 3: VISION MODEL FALLBACK CHAIN (moondream ekle)
    vision_chain = [
       # ✅ En hafif - 1.6B params, çok hızlı (generic için ideal)
        "llava:latest",  # Standard
        "llava",         # Fallback
        "bakllava",      # Son çare
    ]
    
    # Seçilen görselleri analiz et
    for img_data in best_images:
        try:
            logger.info(f"[RAG] 🖼️ Analyzing BEST image: {img_data['safe_path'].name} (Score: {img_data['score']})")
            
            import httpx
            import base64
            
            # Read and encode image
            with open(img_data['safe_path'], "rb") as img_file:
                img_base64 = base64.b64encode(img_file.read()).decode("utf-8")
            
            ollama_url = settings.ollama_base_url
            analysis = None
            
            # ✅ Try vision chain (moondream first for speed)
            for vision_model in vision_chain:
                try:
                    logger.info(f"[RAG] Trying vision model: {vision_model}")
                    response = httpx.post(
                        f"{ollama_url}/api/chat",
                        json={
                            "model": vision_model,
                            "messages": [
                                {
                                    "role": "user",
                                    "content": "Bu görseli detaylı bir şekilde Türkçe olarak açıkla. Görseldeki grafiklerin, tabloların veya şemaların ne anlattığını Türkçe yaz. Sadece Türkçe cevap ver.",
                                    "images": [img_base64]
                                }
                            ],
                            "stream": False,
                            "options": {"num_gpu": -1},
                        },
                        timeout=120.0  # Generic: Enough time for model switching
                    )

                    if response.status_code == 200:
                        response_data = response.json()
                        analysis = response_data.get("message", {}).get("content", "")
                        # ✅✅✅ BU SATIRLARI EKLEYİN ✅✅✅
                        logger.info(f"[VISION RESPONSE] Model: {vision_model}")
                        logger.info(f"[VISION RESPONSE] Status: SUCCESS")
                        logger.info(f"[VISION RESPONSE] Length: {len(analysis)} chars")
                        logger.info(f"[VISION RESPONSE] Full Content:")
                        logger.info("=" * 80)
                        logger.info(analysis)
                        logger.info("=" * 80)

                        if analysis and len(analysis.strip()) > 30:
                            logger.info(f"[RAG] ✅ Vision SUCCESS with {vision_model}")
                            break
                        else:
                            logger.warning(f"[RAG] ⚠️ {vision_model} returned empty/short response")
                    else:
                        # ✅✅✅ HATA DURUMUNDA DA DETAY ✅✅✅
                        logger.error(f"[VISION RESPONSE] Model: {vision_model}")
                        logger.error(f"[VISION RESPONSE] Status: FAILED")
                        logger.error(f"[VISION RESPONSE] HTTP Code: {response.status_code}")
                        logger.error(f"[VISION RESPONSE] Response: {response.text}")
                except Exception as model_err:
                    logger.warning(f"[RAG] ❌ {vision_model} failed: {model_err}")
                    continue
            
            # Honest logging
            vision_success = bool(analysis and len(analysis.strip()) > 30)
            
            if not vision_success:
                analysis = f"Görsel: {img_data['safe_path'].name} (Vision model kullanılamıyor - grafik/tablo/diyagram olabilir)"
                logger.warning(f"[RAG] ⚠️ Vision FAILED - using fallback")
            else:
                logger.info(f"[RAG] ✅ Vision SUCCESS")
            
            image_url = f"{base_url}/{img_data['path'].lstrip('/')}"
            markdown_line = f"![{img_data['alt'] or 'Chart/Graph'}]({image_url})"
            
            image_blocks.append({
                "source": img_data['source'],
                "filename": img_data['safe_path'].name,
                "markdown": markdown_line,
                "analysis": analysis,
                "vision_success": vision_success,
                "relevance_score": img_data['score'],
            })
            
        except Exception as e:
            logger.error(f"[RAG] Error analyzing image: {e}")
    
    # ============ FILTER IRRELEVANT IMAGES FROM CONTEXT ============
    # Remove images that don't match query keywords BEFORE sending to LLM
    processed_docs = []
    base_url = settings.backend_public_url.rstrip("/")
    
    for doc in retrieved_docs:
        content = doc.page_content
        # Replace image paths with full URLs
        content = re.sub(
            r'!\[([^\]]*)\]\(([^)]+)\)',
            lambda m: f'![{m.group(1)}]({base_url}/{m.group(2).lstrip("/")})',
            content
        )
        processed_docs.append((doc.metadata.get('source', 'unknown'), content))

    # Text context
    text_context = "\n\n".join(
        f"📄 Kaynak: {source}\n{content}"
        for source, content in processed_docs
    )

    # ✅ CRITICAL: Build combined context with images
    if image_blocks:
        images_context = "\n\n".join(
            f"🖼️ GÖRSEL ANALİZİ\n"
            f"Kaynak: {b['source']}\n"
            f"Dosya: {b['filename']}\n"
            f"Görsel: {b['markdown']}\n\n"
            f"İÇERİK:\n{b['analysis']}"
            for b in image_blocks
        )
        
        combined_context = (
            "DÖKÜMAN BİLGİLERİ:\n\n"
            "=== METİN ===\n"
            f"{text_context}\n\n"
            "=== GÖRSELLERİN DETAYLI ANALİZİ ===\n"
            f"{images_context}\n\n"
            "⚠️ Hem metin hem görsel bilgilerini kullan!"
        )
        
        logger.info(f"[RAG] ✅ Context ready: {len(retrieved_docs)} chunks + {len(image_blocks)} images")
    else:
        combined_context = (
            "DÖKÜMAN BİLGİLERİ:\n\n"
            f"{text_context}"
        )
        logger.info(f"[RAG] Context ready: {len(retrieved_docs)} chunks (no images)")

    # ✅ DEBUG: Log FULL context
    logger.info(f"[RAG DEBUG] FULL CONTEXT ({len(combined_context)} chars):")
    logger.info("=" * 80)
    logger.info(combined_context[:1000])  # First 1000 chars
    logger.info("..." if len(combined_context) > 1000 else "")
    logger.info("=" * 80)

    return combined_context, retrieved_docs


# ============ RAG AGENT ============

RAG_SYSTEM_PROMPT = """SEN BİR DÖKÜMAN ANALİZ ASİSTANISIN.

🇹🇷 **DİL KURALI**: HER ZAMAN TÜRKÇE CEVAP VER!

📋 **GÖREV**: 
- `retrieve_context` tool'unu kullanarak ilgili bilgileri getir
- Tool'dan gelen BİLGİLERİ KULLANARAK cevap ver
- Tool iki tür bilgi döndürür:
  1. DÖKÜMAN METİN PARÇALARI (yazılı içerik)
  2. GÖRSEL İÇERİK ANALİZLERİ (vision model'in görsel analizi)

⚠️ **ÖNEMLİ KURALLAR**:

1. **Tool Sonuçlarını Kullan**:
    - Tool'dan gelen HEM METİN HEM DE GÖRSEL bilgilerini birleştir
    - "Görselin Detaylı İçeriği" bölümündeki analizi MUTLAKA dahil et
    - Sadece tool'dan gelen bilgileri kullan, kendi bilgini ekleme

2. **Görsel Bilgileri Açıkla**:
    - Tool'dan "🖼️ GÖRSEL İÇERİK ANALİZİ" başlığı altında görsel bilgisi geliyorsa:
    - Bu bilgiyi Türkçe özetle ve kullanıcıya aktar
    - Örnek: "Belgede yer alan görselde [vision analizi]"

3. **Kaynak Belirt**:
    - "Belgede..." veya "PDF'te..." diye başla
    - Hem metin hem görsel kaynaklı bilgileri harmanlayarak ver

4. **Bilgi Yoksa Söyle**:
    - Tool bilgi bulamadıysa: "Bu bilgi belgede bulunmuyor"
    - Tahmin etme, kendi bilgini kullanma

✅ **ÖRNEK CEVAP**:
Kullanıcı: "Yöntemden bahseder misin?"

Doğru Cevap:
"Belgede yöntem başlığı altında şu bilgiler yer alıyor:

[Tool'dan gelen metin bilgisi]

Ayrıca belgede yer alan görselde [vision model'in analizi - örn: "6 aşamalı bir süreç şeması gösterilmektedir: Veri Toplama, Ön İşleme, Model Eğitimi..."]

[Görseli göster]"

Şimdi kullanıcının sorusunu cevapla!
"""

def _get_rag_model():
    """Get LLM model for RAG agent - using Groq (Llama 3.3 70B) for high intelligence"""
    # RAG kritik bir görev -> Groq kullan (güvenilir, hosted, 70B parameter)
    if settings.groq_api_key:
        from langchain_groq import ChatGroq
        logger.info("[RAG] Using Groq (llama-3.3-70b-versatile) for expert reasoning")
        return ChatGroq(
            model="llama-3.3-70b-versatile",
            api_key=settings.groq_api_key,
            temperature=0  # ZERO temperature - no creativity, pure fact extraction
        )
    
    # Fallback: settings default (usually local ollama or gemini)
    logger.info(f"[RAG] Fallback to default model: {settings.default_model}")
    provider, model_name = settings.get_model_provider(settings.default_model)
    return init_chat_model(model=model_name, model_provider=provider, temperature=0)


# Create RAG agent
setup_langsmith()

# CRITICAL: RAG agent should ONLY use retrieve_context
# All file ingestion (PDF, DOCX, etc.) is done BEFORE querying via /api/rag/ingest
# Agent must NOT try to load files during query time
_rag_tools = [
    retrieve_context  # ONLY tool - retrieves pre-processed chunks
]

# Lazy init graph - model created on demand
_graph_instance = None

def _get_graph():
    """Graph creation - NO CHANGES NEEDED"""
    global _graph_instance
    if _graph_instance is not None:
        return _graph_instance
    model = _get_rag_model()
    logger.info(f"[RAG] Using model: {type(model).__name__}")
    # ✅ This works correctly now - tool results flow naturally
    _graph_instance = create_react_agent(
        model, 
        _rag_tools,
        prompt=RAG_SYSTEM_PROMPT  # No {context} placeholder
    )
    return _graph_instance

# Export graph via property with image extraction
class _GraphProxy:
    def __getattr__(self, name):
        # Intercept ainvoke to extract images
        if name == 'ainvoke':
            async def ainvoke_with_images(input_data, *args, **kwargs):
                result = await _get_graph().ainvoke(input_data, *args, **kwargs)
                
                # Extract images from the conversation
                images = []
                if isinstance(result, dict) and "messages" in result:
                    for msg in result["messages"]:
                        if hasattr(msg, "content") and isinstance(msg.content, str):
                            # Find image references
                            import re
                            base_url = settings.backend_public_url.rstrip("/")
                            escaped = re.escape(base_url)
                            img_matches = re.findall(
                                rf'!\[([^\]]*)\]\(({escaped}/[^)]+)\)',
                                msg.content,
                            )
                            for alt, url in img_matches:
                                if not any(img['url'] == url for img in images):
                                    images.append({
                                        "url": url,
                                        "alt": alt or "Chart/Graph",
                                    })
                
                # Add images to result
                if isinstance(result, dict):
                    result["images"] = images
                
                return result
            return ainvoke_with_images
        
        return getattr(_get_graph(), name)

graph = _GraphProxy()


# ============ UTILITY FUNCTIONS ============

def clear_vector_store():
    """Clear all documents from vector store and reset tracking (including Supabase)"""
    global _vector_store, _ingested_sources
    
    # Delete all documents from ChromaDB
    try:
        # Get all IDs and delete them
        collection = _vector_store._collection
        all_ids = collection.get()["ids"]
        if all_ids:
            collection.delete(ids=all_ids)
            logger.info(f"[RAG] Deleted {len(all_ids)} documents from ChromaDB")
    except Exception as e:
        logger.warning(f"[RAG] Failed to clear ChromaDB: {e}")
    
    _ingested_sources.clear()
    
    # Clear Supabase records (if configured)
    try:
        if settings.supabase_url and settings.supabase_key:
            try:
                from supabase import create_client
            except Exception:
                create_client = None
            if create_client:
                client = create_client(settings.supabase_url, settings.supabase_key)
                # Delete all records from rag_sources table
                client.table("rag_sources").delete().neq("source_name", "").execute()
                logger.info("[RAG] 🗑️ Supabase rag_sources table cleared")
    except Exception as e:
        logger.warning(f"[RAG] Failed to clear Supabase rag_sources: {e}")
    
    # Persist empty state to disk
    _persist_ingested_sources()
    
    logger.info("[RAG] 🗑️ Vector store cleared (memory + disk + Supabase)")


def remove_source(source_name: str) -> bool:
    """
    Remove all chunks for a specific source from vector store.
    
    Args:
        source_name: Source identifier to remove
        
    Returns:
        True if source was found and removed
    """
    if source_name not in _ingested_sources:
        logger.warning(f"[RAG] ⚠️ Source not found: {source_name}")
        return False
    
    # 🚀 PERFORMANCE: Use ChromaDB as source of truth (no in-memory list)
    collection = _vector_store._collection
    all_data = collection.get()
    
    ids_to_delete = [
        doc_id for doc_id, meta in zip(all_data['ids'], all_data['metadatas'])
        if meta.get('source') == source_name
    ]
    
    if ids_to_delete:
        collection.delete(ids=ids_to_delete)
        logger.info(f"[RAG] 🗑️ Deleted {len(ids_to_delete)} chunks from ChromaDB")
    
    # Remove from tracking
    _ingested_sources.discard(source_name)
    _persist_ingested_sources()

    logger.info(f"[RAG] 🗑️ Removed source: {source_name}")
    return True


def get_ingested_sources() -> list:
    """Get list of all ingested source names"""
    return list(_ingested_sources)


# Load persisted ingested sources at import time (best-effort)
_load_ingested_sources()
