"""
RAG Agent with LangChain - Document Processing & Semantic Search
Based on: https://docs.langchain.com/oss/python/langchain/rag

Features:
- PDF, DOCX, Audio, Image processing
- Vector store (Supabase pgvector)
- Semantic search with embeddings
- LangGraph agent with retrieval tool
"""

import os
import logging
from typing import List, Optional, Literal, Union
from pathlib import Path

from langchain.chat_models import init_chat_model
from langchain.tools import tool
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document
from langgraph.prebuilt import create_react_agent
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_ollama import OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
try:
    from langchain_community.document_loaders import (
        PyPDFLoader,
        Docx2txtLoader,
        CSVLoader,
    )
except ImportError:
    # Fallback: manual loaders
    PyPDFLoader = None
    Docx2txtLoader = None
    CSVLoader = None

import pymupdf4llm  # High-quality PDF to Markdown

from ..config import settings

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ============ LANGSMITH ============
def setup_langsmith():
    """LangSmith tracing for RAG agent"""
    tracing_enabled = os.getenv("LANGSMITH_TRACING", "").strip().lower() in {"1", "true", "yes", "on"}
    if not tracing_enabled:
        return False
    if not settings.langsmith_api_key:
        return False

    os.environ["LANGSMITH_TRACING"] = "true"
    os.environ["LANGSMITH_API_KEY"] = settings.langsmith_api_key

    # Standard LangChain vars
    os.environ["LANGCHAIN_TRACING_V2"] = "true"
    os.environ["LANGCHAIN_API_KEY"] = settings.langsmith_api_key
    os.environ["LANGCHAIN_PROJECT"] = os.getenv("LANGCHAIN_PROJECT", "ai-research-rag-agent")

    logger.info("[LANGSMITH] Tracing aktif")
    return True


# ============ VECTOR STORE SETUP ============

# Initialize embeddings (Ollama - Local & Free)
# Override via env: OLLAMA_EMBED_MODEL
_embeddings = OllamaEmbeddings(
    model=os.getenv("OLLAMA_EMBED_MODEL", "nomic-embed-text:latest"),
    base_url=settings.ollama_base_url,
)

# Initialize vector store (in-memory for now, Supabase later)
_vector_store = InMemoryVectorStore(_embeddings)

# Track ingested sources to prevent duplicates
_ingested_sources = set()

# Keep an in-memory registry of stored chunks so we can rebuild the vector store
# when deleting by source (InMemoryVectorStore has no metadata delete).
_stored_chunks: List[Document] = []


def _uploads_dir() -> Path:
    # Normalize to project uploads directory
    return Path(settings.project_root) / "uploads"


_INGESTED_SOURCES_FILE = _uploads_dir() / ".rag_ingested_sources.json"


def _load_ingested_sources() -> None:
    """Load ingested sources from disk (and optionally Supabase if configured).

    This is best-effort; it should never crash import.
    """
    global _ingested_sources

    # 1) Disk persistence (always available)
    try:
        if _INGESTED_SOURCES_FILE.exists():
            import json
            data = json.loads(_INGESTED_SOURCES_FILE.read_text(encoding="utf-8"))
            if isinstance(data, list):
                _ingested_sources.update(str(x) for x in data if str(x).strip())
    except Exception as e:
        logger.warning(f"[RAG] Failed to load ingested sources from disk: {e}")

    # 2) Supabase persistence (optional)
    # Requires a table named `rag_sources` with at least:
    # - source_name text primary key
    # - created_at timestamptz default now()
    try:
        if settings.supabase_url and settings.supabase_key:
            try:
                from supabase import create_client  # type: ignore
            except Exception:
                create_client = None
            if create_client:
                client = create_client(settings.supabase_url, settings.supabase_key)
                resp = client.table("rag_sources").select("source_name").limit(5000).execute()
                for row in resp.data or []:
                    name = row.get("source_name")
                    if name:
                        _ingested_sources.add(str(name))
    except Exception as e:
        # Don't fail if table doesn't exist or credentials are limited.
        logger.info(f"[RAG] Supabase rag_sources not loaded (optional): {e}")


def _persist_ingested_sources() -> None:
    """Persist ingested sources to disk and (optionally) Supabase."""
    # 1) Disk
    try:
        import json
        _uploads_dir().mkdir(parents=True, exist_ok=True)
        _INGESTED_SOURCES_FILE.write_text(
            json.dumps(sorted(_ingested_sources), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception as e:
        logger.warning(f"[RAG] Failed to persist ingested sources to disk: {e}")

    # 2) Supabase (best-effort)
    try:
        if settings.supabase_url and settings.supabase_key:
            try:
                from supabase import create_client  # type: ignore
            except Exception:
                create_client = None
            if create_client:
                client = create_client(settings.supabase_url, settings.supabase_key)
                # Upsert all sources (requires primary key on source_name)
                rows = [{"source_name": s} for s in _ingested_sources]
                if rows:
                    client.table("rag_sources").upsert(rows).execute()
    except Exception as e:
        logger.info(f"[RAG] Supabase rag_sources not persisted (optional): {e}")

# ============ GLOBAL CHUNKING STRATEGY (ROBUST) ============
# Goal (global): keep "image + its local explanation" together, and avoid micro-chunks.
# Why not SemanticChunker here:
# - For PDF→Markdown, it often splits into single-sentence chunks (bad retrieval and bad answers)
# - It can cut across headings/images, mixing sections
#
# Strategy:
# 1) If content looks like Markdown (headings/images), split by headings up to H4.
#    (We intentionally DO NOT hard-split on H5/H6 to avoid separating algorithm heading from its description.)
# 2) Inside each section, split by images so each image gets its own chunk when multiple exist.
# 3) If a chunk is still too large, fall back to a gentle recursive splitter.

import re


def _looks_like_markdown(text: str) -> bool:
    if not text:
        return False
    # headings or markdown images are strong signals
    return ("\n#" in text) or ("![](" in text) or bool(re.search(r"(?m)^#{1,6}\s+", text))


_fallback_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1200,
    chunk_overlap=120,
    add_start_index=True,
    separators=[
        "\n\n",  # paragraphs
        "\n",    # lines
        " ",
        "",
    ],
    is_separator_regex=False,
)


def _split_markdown_by_headings(md: str, max_heading_level: int = 4) -> List[str]:
    """Split markdown into sections using headings up to max_heading_level."""
    md = (md or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = md.split("\n")
    sections: List[List[str]] = []
    current: List[str] = []

    heading_re = re.compile(r"^(#{1,6})\s+.*")
    for line in lines:
        m = heading_re.match(line.strip())
        if m:
            level = len(m.group(1))
            if level <= max_heading_level and current:
                sections.append(current)
                current = [line]
                continue
        current.append(line)

    if current:
        sections.append(current)

    return ["\n".join(s).strip() for s in sections if "\n".join(s).strip()]


def _split_section_by_images(section_text: str) -> List[str]:
    """Within a section, split into smaller chunks by image boundaries.

    If there are multiple images in the same section (e.g., conclusion page),
    each image becomes its own chunk along with nearby text.
    """
    section_text = (section_text or "").strip()
    if not section_text:
        return []

    lines = section_text.split("\n")
    image_line_re = re.compile(r"^!\[[^\]]*\]\([^)]+\)\s*$")

    chunks: List[str] = []
    current: List[str] = []
    current_has_image = False
    recent_nonempty: List[str] = []

    for line in lines:
        is_image = bool(image_line_re.match(line.strip()))

        if is_image and current_has_image:
            # finalize previous image-chunk
            chunk = "\n".join(current).strip()
            if chunk:
                chunks.append(chunk)
            current = []
            current_has_image = False

        if is_image and not current:
            # pull a little context before the image, if available
            for ctx in recent_nonempty[-6:]:
                current.append(ctx)

        current.append(line)
        if is_image:
            current_has_image = True

        if line.strip():
            recent_nonempty.append(line)

    tail = "\n".join(current).strip()
    if tail:
        chunks.append(tail)

    return chunks


def _make_rag_splits(text: str, source_name: str) -> List[Document]:
    """Create RAG-ready Document splits with stable semantics."""
    base_doc = Document(page_content=text or "", metadata={"source": source_name})

    if not _looks_like_markdown(base_doc.page_content):
        splits = _fallback_splitter.split_documents([base_doc])
        for s in splits:
            s.metadata["source"] = source_name
            s.metadata["has_images"] = False
        return splits

    # Markdown path
    heading_sections = _split_markdown_by_headings(base_doc.page_content, max_heading_level=4)

    chunks: List[Document] = []
    for sec in heading_sections:
        # If section has multiple images, split them into separate chunks.
        subchunks = _split_section_by_images(sec) if "![](" in sec or "![" in sec else [sec]

        for sub in subchunks:
            if not sub.strip():
                continue

            doc = Document(page_content=sub.strip(), metadata={"source": source_name})

            # If a chunk is still too big (rare), gently split by paragraphs.
            if len(doc.page_content) > 2000:
                for s in _fallback_splitter.split_documents([doc]):
                    s.metadata["source"] = source_name
                    s.metadata["has_images"] = bool(re.search(r"!\[[^\]]*\]\([^)]+\)", s.page_content))
                    chunks.append(s)
            else:
                doc.metadata["has_images"] = bool(re.search(r"!\[[^\]]*\]\([^)]+\)", doc.page_content))
                chunks.append(doc)

    return chunks


def _safe_resolve_upload_path(img_path: str) -> Optional[Path]:
    """Resolve a markdown image path safely under uploads/.

    Prevents path traversal (e.g. ../../secret) and disallows absolute paths.
    Returns a resolved absolute Path if valid and exists, else None.
    """
    if not img_path:
        return None

    # Disallow URLs and absolute paths
    lowered = img_path.strip().lower()
    if lowered.startswith("http://") or lowered.startswith("https://"):
        return None

    p = Path(img_path)
    if p.is_absolute():
        return None

    uploads = _uploads_dir().resolve()
    resolved = (uploads / p).resolve()
    try:
        resolved.relative_to(uploads)
    except Exception:
        return None

    if not resolved.exists():
        return None

    return resolved


# ============ DOCUMENT PROCESSING TOOLS ============

@tool
def load_pdf(file_path: str) -> str:
    """
    Load and extract text from a PDF file with images using PyMuPDF4LLM.
    Extracts text, tables, and saves page images.
    
    Args:
        file_path: Absolute path to PDF file
        
    Returns:
        Extracted text content (Markdown)
    """
    # 1. Try PyMuPDF4LLM (Fast, extracts images, good Markdown)
    try:
        import pymupdf4llm
        
        logger.info(f"[PDF] Processing {Path(file_path).name} with PyMuPDF4LLM...")
        
        # Create images folder
        images_folder = Path(file_path).parent / f"{Path(file_path).stem}_images"
        images_folder.mkdir(exist_ok=True)
        
        # Convert PDF to Markdown with images
        md_text = pymupdf4llm.to_markdown(
            file_path,
            write_images=True,
            image_path=str(images_folder),
            image_format="png"
        )
        
        logger.info(f"[PDF] ✅ Converted {len(md_text)} chars with images (PyMuPDF4LLM)")
        
        # Save debug file
        debug_path = Path(file_path).parent / f"{Path(file_path).stem}_pymupdf_debug.md"
        debug_path.write_text(md_text, encoding="utf-8")
        logger.info(f"[PDF] 📝 Markdown saved to: {debug_path}")
        logger.info(f"[PDF] 🖼️ Images saved to: {images_folder}")
        
        if len(md_text) >= 100:
            return md_text
            
    except Exception as e:
        logger.warning(f"[PDF] PyMuPDF4LLM failed ({e}), falling back to Unstructured...")
    
    # 2. Fallback: Unstructured
    try:
        from unstructured.partition.pdf import partition_pdf
        
        logger.info(f"[PDF] Processing {Path(file_path).name} with Unstructured (fallback)...")
        
        elements = partition_pdf(filename=file_path, strategy="fast")
        
        text_parts = []
        for el in elements:
            element_text = str(el)
            element_type = el.category if hasattr(el, 'category') else 'Unknown'
            
            if element_type == 'Title':
                text_parts.append(f"# {element_text}")
            elif element_type == 'Header':
                text_parts.append(f"## {element_text}")
            elif element_type == 'Table':
                text_parts.append(f"\n{element_text}\n")
            else:
                text_parts.append(element_text)
        
        text = "\n\n".join(text_parts)
        logger.info(f"[PDF] ✅ Converted {len(text)} chars from {len(elements)} elements (Unstructured OCR)")
        
        debug_path = Path(file_path).parent / f"{Path(file_path).stem}_unstructured_debug.md"
        debug_path.write_text(text, encoding="utf-8")
        logger.info(f"[PDF] 📝 Unstructured output saved to: {debug_path}")
        
        if len(text) >= 50:
            return text
            
    except Exception as e2:
        logger.warning(f"[PDF] Unstructured failed ({e2}), falling back to PyPDF...")

    # 3. Last Fallback: PyPDF (Simple text extraction)
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = "\n\n".join([page.extract_text() for page in reader.pages])
        logger.info(f"[PDF] Loaded {len(reader.pages)} pages, {len(text)} chars (PyPDF - last resort)")
        
        debug_path = Path(file_path).parent / f"{Path(file_path).stem}_pypdf_debug.txt"
        debug_path.write_text(text, encoding="utf-8")
        logger.info(f"[PDF] 📝 PyPDF output saved to: {debug_path}")
        
        return text
    except Exception as e3:
        logger.error(f"[PDF] All methods failed: {e3}")
        return f"Error loading PDF: {str(e3)}"


@tool
def load_docx(file_path: str) -> str:
    """
    Load and extract text + images from a DOCX file using python-docx.
    Extracts text, tables, and saves embedded images similar to PDF processing.
    
    Args:
        file_path: Absolute path to DOCX file
        
    Returns:
        Extracted text content (Markdown with image references)
    """
    # 1. Try python-docx with image extraction (Primary method)
    try:
        from docx import Document as DocxDocument
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        import io
        
        logger.info(f"[DOCX] Processing {Path(file_path).name} with python-docx...")
        
        # Create images folder
        images_folder = Path(file_path).parent / f"{Path(file_path).stem}_images"
        images_folder.mkdir(exist_ok=True)
        
        doc = DocxDocument(file_path)
        markdown_parts = []
        image_counter = 0
        
        # Extract images from document
        image_map = {}  # Map relationship ID to saved filename
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    # Determine extension from content type
                    ext_map = {
                        'image/png': 'png',
                        'image/jpeg': 'jpg',
                        'image/jpg': 'jpg',
                        'image/gif': 'gif',
                        'image/bmp': 'bmp'
                    }
                    ext = ext_map.get(rel.target_part.content_type, 'png')
                    
                    # Save image
                    image_filename = f"{Path(file_path).stem}-{image_counter}.{ext}"
                    image_path = images_folder / image_filename
                    image_path.write_bytes(image_data)
                    
                    # Store relative path for markdown
                    relative_path = f"{images_folder.name}/{image_filename}"
                    image_map[rel_id] = relative_path
                    image_counter += 1
                    
                except Exception as e:
                    logger.warning(f"[DOCX] Failed to extract image {rel_id}: {e}")
        
        # Process paragraphs and tables in order
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                text = para.text.strip()
                
                # Check paragraph style for headers
                style_name = para.style.name if para.style else ""
                
                if style_name.startswith('Heading 1'):
                    markdown_parts.append(f"# {text}")
                elif style_name.startswith('Heading 2'):
                    markdown_parts.append(f"## {text}")
                elif style_name.startswith('Heading 3'):
                    markdown_parts.append(f"### {text}")
                elif style_name.startswith('Heading 4'):
                    markdown_parts.append(f"#### {text}")
                elif text:
                    markdown_parts.append(text)
                
                # Check for inline images
                for run in para.runs:
                    for rel_id in run.element.xpath('.//a:blip/@r:embed', namespaces={
                        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                    }):
                        if rel_id in image_map:
                            markdown_parts.append(f"![Image]({image_map[rel_id]})")
                
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                # Convert table to markdown
                markdown_parts.append("\n")
                for i, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    markdown_parts.append("| " + " | ".join(cells) + " |")
                    if i == 0:  # Add separator after header
                        markdown_parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
                markdown_parts.append("\n")
        
        md_text = "\n\n".join(markdown_parts)
        logger.info(f"[DOCX] ✅ Converted {len(md_text)} chars with {image_counter} images (python-docx)")
        
        # Save debug file
        debug_path = Path(file_path).parent / f"{Path(file_path).stem}_docx_debug.md"
        debug_path.write_text(md_text, encoding="utf-8")
        logger.info(f"[DOCX] 📝 Markdown saved to: {debug_path}")
        if image_counter > 0:
            logger.info(f"[DOCX] 🖼️ {image_counter} images saved to: {images_folder}")
        
        if len(md_text) >= 50:
            return md_text
            
    except Exception as e:
        logger.warning(f"[DOCX] python-docx failed ({e}), falling back to Unstructured...")
    
    # 2. Fallback: Unstructured (no image extraction)
    try:
        from unstructured.partition.docx import partition_docx
        
        logger.info(f"[DOCX] Processing {Path(file_path).name} with Unstructured (fallback)...")
        
        elements = partition_docx(filename=file_path)
        text = "\n\n".join([str(el) for el in elements])
        
        logger.info(f"[DOCX] ✅ Converted {len(text)} chars (Unstructured)")
        
        debug_path = Path(file_path).parent / f"{Path(file_path).stem}_unstructured_debug.md"
        debug_path.write_text(text, encoding="utf-8")
        logger.info(f"[DOCX] 📝 Markdown saved to: {debug_path}")
        
        return text
        
    except Exception as e2:
        logger.warning(f"[DOCX] Unstructured failed ({e2}), falling back to Docx2txt...")
        
        # 3. Last Fallback: Docx2txt (text only)
        try:
            if Docx2txtLoader:
                loader = Docx2txtLoader(file_path)
                docs = loader.load()
                text = "\n\n".join([doc.page_content for doc in docs])
                return text
            else:
                import docx2txt
                return docx2txt.process(file_path)
        except Exception as e3:
            logger.error(f"[DOCX] All methods failed: {e3}")
            return f"Error loading DOCX: {str(e)}"


@tool
def load_csv(file_path: str) -> str:
    """
    Load and extract data from a CSV file.
    
    Args:
        file_path: Absolute path to CSV file
        
    Returns:
        CSV data as formatted text
    """
    try:
        if CSVLoader:
            loader = CSVLoader(file_path)
            docs = loader.load()
            text = "\n\n".join([doc.page_content for doc in docs])
            logger.info(f"[CSV] Loaded {len(docs)} rows from {Path(file_path).name}")
            return text
        else:
            # Fallback: csv module
            import csv
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                rows = [str(dict(row)) for row in reader]
                text = "\n\n".join(rows)
                logger.info(f"[CSV] Loaded {len(rows)} rows from {Path(file_path).name}")
                return text
    except Exception as e:
        logger.error(f"[CSV] Error loading {file_path}: {e}")
        return f"Error loading CSV: {str(e)}"


@tool
async def transcribe_audio(file_path: str, language: str = "tr") -> str:
    """
    Transcribe audio file locally.
    
    Args:
        file_path: Absolute path to audio file (mp3, wav, m4a)
        language: Language code (default: tr for Turkish)
        
    Returns:
        Transcribed text
    """
    try:
        # Local STT recommendation: faster-whisper
        # Install (optional): pip install faster-whisper ffmpeg-python
        try:
            from faster_whisper import WhisperModel  # type: ignore
        except Exception:
            return (
                "Local STT için `faster-whisper` önerilir.\n"
                "Kurulum: `pip install faster-whisper ffmpeg-python` ve sistemde ffmpeg olmalı.\n"
                "Alternatif: whisper.cpp (CLI)."
            )

        model_name = os.getenv("WHISPER_MODEL", "small")
        device = os.getenv("WHISPER_DEVICE", "cpu")
        compute_type = os.getenv("WHISPER_COMPUTE_TYPE", "int8")

        whisper_model = WhisperModel(model_name, device=device, compute_type=compute_type)
        segments, info = whisper_model.transcribe(file_path, language=language)
        text = "".join(seg.text for seg in segments).strip()
        logger.info(f"[AUDIO] Transcribed {Path(file_path).name} ({info.language}) - {len(text)} chars")
        return text or "(Boş transkripsiyon)"
    except Exception as e:
        logger.error(f"[AUDIO] Error transcribing {file_path}: {e}")
        return f"Error transcribing audio: {str(e)}"


@tool
async def analyze_image(file_path: str, prompt: str = "Bu görseli detaylı Türkçe açıkla") -> str:
    """
    Analyze image locally via Ollama vision model.
    
    Args:
        file_path: Absolute path to image file
        prompt: Analysis prompt (default: Turkish description)
        
    Returns:
        Image analysis text
    """
    try:
        import base64
        import httpx

        with open(file_path, "rb") as image_file:
            image_b64 = base64.b64encode(image_file.read()).decode("utf-8")

        vision_model = os.getenv("OLLAMA_VISION_MODEL", "llava:latest")

        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                f"{settings.ollama_base_url}/api/chat",
                json={
                    "model": vision_model,
                    "messages": [
                        {"role": "user", "content": prompt, "images": [image_b64]}
                    ],
                    "stream": False,
                },
            )

        if response.status_code != 200:
            return f"Ollama Vision error: HTTP {response.status_code} - {response.text}"

        data = response.json()
        text = (data.get("message", {}) or {}).get("content", "")
        logger.info(f"[IMAGE] Analyzed {Path(file_path).name} with {vision_model}")
        return text or "(Boş analiz)"
    except Exception as e:
        logger.error(f"[IMAGE] Error analyzing {file_path}: {e}")
        return f"Error analyzing image: {str(e)}"


# ============ RAG CORE FUNCTIONS ============

def ingest_documents(documents: List[Document], source_name: str = "upload") -> int:
    """
    Ingest documents into vector store.
    
    Args:
        documents: List of Document objects
        source_name: Source identifier for metadata
        
    Returns:
        Number of chunks added
    """
    # Split documents into chunks (robust, image-aware)
    all_splits: List[Document] = []
    for doc in documents:
        all_splits.extend(_make_rag_splits(doc.page_content, source_name))
    
    # Add source metadata
    for split in all_splits:
        split.metadata["source"] = source_name
    
    # Add to vector store + registry
    if all_splits:
        _vector_store.add_documents(documents=all_splits)
        _stored_chunks.extend(all_splits)
        _ingested_sources.add(source_name)
        _persist_ingested_sources()
    
    logger.info(f"[RAG] Ingested {len(all_splits)} chunks from {source_name}")
    return len(all_splits)


def ingest_text(text: str, source_name: str = "upload") -> int:
    """
    Ingest raw text into vector store with intelligent chunking.
    
    Supports two modes (set via USE_AGENTIC_CHUNKER env var):
    - Agentic Chunking (default): LLM-powered semantic grouping
    - Markdown Splitting (fallback): Header-based splitting
    
    Args:
        text: Raw text content (Markdown preferred)
        source_name: Source identifier
        
    Returns:
        Number of chunks added (0 if already exists)
    """
    # Check if already ingested
    if source_name in _ingested_sources:
        logger.warning(f"[RAG] ⚠️ Duplicate: {source_name} already ingested, skipping...")
        return 0
    
    # Choose chunking strategy based on environment
    use_agentic = os.getenv("USE_AGENTIC_CHUNKER", "true").lower() in ("1", "true", "yes", "on")
    
    if use_agentic:
        # Agentic Chunking: LLM decides which content belongs together
        try:
            from .agentic_chunker import agentic_chunk_text
            logger.info(f"[RAG] 🤖 Using Agentic Chunker (Gemini 2.5) for: {source_name}")
            final_splits = agentic_chunk_text(text, source_name)
        except Exception as e:
            logger.warning(f"[RAG] Agentic chunking failed ({e}), falling back to markdown splitter")
            final_splits = _make_rag_splits(text, source_name)
    else:
        # Traditional: Markdown header-based splitting
        logger.info(f"[RAG] 📄 Using Markdown splitter for: {source_name}")
        final_splits = _make_rag_splits(text, source_name)

    # Add source metadata
    for split in final_splits:
        split.metadata["source"] = source_name

    # Add to vector store + registry
    if final_splits:
        _vector_store.add_documents(documents=final_splits)
        _stored_chunks.extend(final_splits)
        _ingested_sources.add(source_name)  # Mark as ingested
        _persist_ingested_sources()
    
    logger.info(f"[RAG] ✅ Ingested {len(final_splits)} chunks from {source_name}")
    return len(final_splits)


# ============ RETRIEVAL TOOL ============

from pydantic import BaseModel, Field

class RetrieveContextInput(BaseModel):
    """Input schema for retrieve_context tool"""
    query: str = Field(description="Search query to find relevant document chunks")
    # Groq API strict validation fix: Accept string, convert to int internally
    top_k: str = Field(default="3", description="Number of top results to return (e.g. '3')")

@tool(args_schema=RetrieveContextInput, response_format="content_and_artifact")
def retrieve_context(query: str, top_k: str = "3"):
    """
    Retrieve relevant context from ingested documents using semantic search.
    AUTOMATICALLY analyzes images found in the context using vision models.
    
    Args:
        query: Search query
        top_k: Number of top results to return (default: "3")
        
    Returns:
        Retrieved context as formatted text with source metadata AND image analysis
    """
    # Ensure top_k is an integer
    try:
        k_val = int(top_k)
    except (ValueError, TypeError):
        k_val = 3

    retrieved_docs = _vector_store.similarity_search(query, k=k_val)
    
    if not retrieved_docs:
        return "No relevant documents found.", []
    
    # Extract image paths from markdown content
    import re
    image_blocks = []
    
    base_url = settings.backend_public_url.rstrip("/")
    for doc in retrieved_docs:
        # Find all markdown image references: ![alt](path)
        img_matches = re.findall(r'!\[([^\]]*)\]\(([^)]+)\)', doc.page_content)

        for alt_text, img_path in img_matches:
            safe_path = _safe_resolve_upload_path(img_path)
            if not safe_path:
                continue

            try:
                logger.info(f"[RAG] 🖼️ Analyzing image: {safe_path.name}")
                analysis = analyze_image.invoke(str(safe_path))

                image_url = f"{base_url}/{img_path.lstrip('/')}"
                markdown_line = f"![{alt_text or 'Chart/Graph'}]({image_url})"

                image_blocks.append({
                    "source": doc.metadata.get("source", "unknown"),
                    "filename": safe_path.name,
                    "markdown": markdown_line,
                    "analysis": analysis,
                })
            except Exception as e:
                logger.error(f"[RAG] Failed to analyze image {safe_path}: {e}")
    
    # Replace image paths in chunks with full URLs BEFORE sending to LLM
    processed_docs = []
    for doc in retrieved_docs:
        content = doc.page_content
        # Replace all image markdown with full URLs
        base_url = settings.backend_public_url.rstrip("/")
        content = re.sub(
            r'!\[([^\]]*)\]\(([^)]+)\)',
            lambda m: f'![{m.group(1)}]({base_url}/{m.group(2).lstrip("/")})',
            content
        )
        processed_docs.append((doc.metadata.get('source', 'unknown'), content))
    
    # Build context
    text_context = "\n\n".join(
        (f"📄 Source: {source}\n{content}")
        for source, content in processed_docs
    )

    # IMPORTANT: Provide images as individual blocks (image -> analysis), not a single combined list.
    if image_blocks:
        images_context = "\n\n".join(
            (
                "🖼️ IMAGE BLOCK\n"
                f"Source: {b['source']}\n"
                f"File: {b['filename']}\n"
                f"Markdown: {b['markdown']}\n"
                "Vision analysis (use this to write your own caption in Turkish):\n"
                f"{b['analysis']}"
            )
            for b in image_blocks
        )
        combined_context = text_context + "\n\n" + images_context
        logger.info(f"[RAG] Retrieved {len(retrieved_docs)} chunks + {len(image_blocks)} image blocks")
    else:
        combined_context = text_context
        logger.info(f"[RAG] Retrieved {len(retrieved_docs)} chunks (no images found)")
    
    return combined_context, retrieved_docs


# ============ RAG AGENT ============

RAG_SYSTEM_PROMPT = """You are a RAG agent with vision. Documents are pre-loaded.

WORKFLOW:
1. Call retrieve_context with user's query keywords
2. Analyze returned text + image analyses  
3. Generate answer using retrieved information

IMAGE HANDLING:
- When you see an IMAGE BLOCK, you MUST present it as:
    1) the markdown image line
    2) immediately below, 1-2 sentences of your OWN explanation in Turkish
- Do NOT output a single combined list of all image captions.
- Keep each image + its explanation together.

TOOL: retrieve_context(query, top_k="3")
- Use for all document questions
- Automatically analyzes images in chunks
- Returns text + vision analysis

RESPONSE RULES:
- Use markdown formatting
 - When context includes an IMAGE BLOCK with "Markdown:", copy that markdown line into your answer
 - Then write a short Turkish caption (your own sentence), using the vision analysis + nearby text
 - If multiple images are relevant, output them one by one (image then caption)
"""


def _get_rag_model():
    """Get LLM model for RAG agent - use Groq for critical RAG tasks"""
    # RAG kritik bir görev -> Groq kullan (güvenilir, hosted)
    if settings.groq_api_key:
        from langchain_groq import ChatGroq
        logger.info("[RAG] Using Groq (llama-3.3-70b-versatile) for critical RAG tasks")
        return ChatGroq(
            model="llama-3.3-70b-versatile",
            api_key=settings.groq_api_key,
            temperature=0
        )
    # Fallback: settings default
    logger.info(f"[RAG] Fallback to default model: {settings.default_model}")
    provider, model_name = settings.get_model_provider(settings.default_model)
    return init_chat_model(model=model_name, model_provider=provider, temperature=0)


# Create RAG agent
setup_langsmith()

_rag_tools = [
    load_pdf,
    load_docx,
    load_csv,
    transcribe_audio,
    analyze_image,
    retrieve_context
]

# Lazy init graph - model created on demand
_graph_instance = None

def _get_graph():
    """Lazy load RAG graph with fresh model"""
    global _graph_instance
    if _graph_instance is not None:
        return _graph_instance
    
    model = _get_rag_model()
    logger.info(f"[RAG] Using model: {type(model).__name__}")
    
    # Use prompt parameter for system message (LangGraph 0.2+)
    _graph_instance = create_react_agent(
        model, 
        _rag_tools,
        prompt=RAG_SYSTEM_PROMPT
    )
    
    return _graph_instance

# Export graph via property with image extraction
class _GraphProxy:
    def __getattr__(self, name):
        # Intercept ainvoke to extract images
        if name == 'ainvoke':
            async def ainvoke_with_images(input_data, *args, **kwargs):
                result = await _get_graph().ainvoke(input_data, *args, **kwargs)
                
                # Extract images from the conversation
                images = []
                if isinstance(result, dict) and "messages" in result:
                    for msg in result["messages"]:
                        if hasattr(msg, "content") and isinstance(msg.content, str):
                            # Find image references
                            import re
                            base_url = settings.backend_public_url.rstrip("/")
                            escaped = re.escape(base_url)
                            img_matches = re.findall(
                                rf'!\[([^\]]*)\]\(({escaped}/[^)]+)\)',
                                msg.content,
                            )
                            for alt, url in img_matches:
                                if not any(img['url'] == url for img in images):
                                    images.append({
                                        "url": url,
                                        "alt": alt or "Chart/Graph",
                                    })
                
                # Add images to result
                if isinstance(result, dict):
                    result["images"] = images
                
                return result
            return ainvoke_with_images
        
        return getattr(_get_graph(), name)

graph = _GraphProxy()


# ============ UTILITY FUNCTIONS ============

def clear_vector_store():
    """Clear all documents from vector store and reset tracking"""
    global _vector_store, _ingested_sources, _stored_chunks
    
    # Re-initialize vector store (clears all data)
    _vector_store = InMemoryVectorStore(_embeddings)
    _ingested_sources.clear()
    _stored_chunks.clear()
    _persist_ingested_sources()
    
    logger.info("[RAG] 🗑️ Vector store cleared")


def remove_source(source_name: str) -> bool:
    """
    Remove all chunks for a specific source from vector store.
    
    Args:
        source_name: Source identifier to remove
        
    Returns:
        True if source was found and removed
    """
    if source_name not in _ingested_sources:
        logger.warning(f"[RAG] ⚠️ Source not found: {source_name}")
        return False
    
    # Rebuild vector store without this source
    remaining = [d for d in _stored_chunks if d.metadata.get("source") != source_name]

    # If registry didn't have anything, still drop from ingested set
    _stored_chunks.clear()
    _stored_chunks.extend(remaining)

    # Re-init and re-add remaining chunks
    global _vector_store
    _vector_store = InMemoryVectorStore(_embeddings)
    if remaining:
        _vector_store.add_documents(documents=remaining)

    _ingested_sources.discard(source_name)
    _persist_ingested_sources()

    logger.info(f"[RAG] 🗑️ Removed source and rebuilt store: {source_name}")
    return True


def get_ingested_sources() -> list:
    """Get list of all ingested source names"""
    return list(_ingested_sources)


# Load persisted ingested sources at import time (best-effort)
_load_ingested_sources()
